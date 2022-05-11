from flask import Flask
import docx
import logging
import uuid
from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient, __version__
app = Flask(__name__)

@app.route("/")
def hello_world():
    
       
    doc = docx.Document()
    doc.add_paragraph('first item in ordered list', style='List Number')
    doc.save('demo.docx')
            
    logging.info('Python HTTP trigger function processed a request.')
    # Create the BlobServiceClient object which will be used to create a container client
    connect_str = "DefaultEndpointsProtocol=https;AccountName=pvpstorage1;AccountKey=tTqnRrNj2jLyxkP3sdYaGraaNWqofISf8SQb71RKC4sSnOmYk0rPcDm0xfvAwpffSqOZ9SVaJ+TO+AStQ1rZnA==;EndpointSuffix=core.windows.net" 
    blob_service_client = BlobServiceClient.from_connection_string(connect_str)
    logging.info("connection Established")
    # Create a unique name for the container
    container_name = str(uuid.uuid4())
    #print(container_name)
            
    #doc_para = doc.add_paragraph('Your paragraph goes here, ')
    logging.info("Connection establised")
    
    # Create the container
    container_client = blob_service_client.create_container(container_name)
    blob_client = blob_service_client.get_blob_client(container=container_name, blob="demo.docx")


    with open("demo.docx", "rb") as data:
        blob_client.upload_blob(data)
    logging.info("uploaded the blob")

    
    
    return "<p>Hello, World!!!!</p>"


















'''
import docx

a= 3
b= 5
print(a+b)
print("success")
doc = docx.Document()
doc.add_paragraph(
    'first item in ordered list pvp' ,
)

doc.save('demo.docx')

doc = docx.Document('demo.docx')
for docpara in doc.paragraphs:
    print(docpara.text)


'''












from docx import Document
document=Document()

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)


document.save("hello.docx")
